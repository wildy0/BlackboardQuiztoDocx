import os
import re

from bs4 import BeautifulSoup, NavigableString
from docx.shared import RGBColor, Mm
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from lxml import etree
import random

new_number = 0
root = None


def generate_pseudo_random_durableId(prefix):
    prefix_str = str(prefix)
    prefix_length = len(prefix_str)

    # Ensure the prefix is not longer than 10 digits
    if prefix_length > 10:
        raise ValueError("The prefix should not be longer than 10 digits.")

    # Calculate the minimum and maximum length for the random part
    min_length = max(1, 9 - prefix_length)
    max_length = 10 - prefix_length

    # Generate the random part
    random_part = random.randint(10 ** (min_length - 1), 10 ** max_length - 1)

    # Concatenate the prefix and the random part
    durableId = int(f"{prefix_str}{random_part}")

    return durableId


def manual_deepcopy(element):
    # Create a new element with the same tag, attributes and text
    new_element = etree.Element(element.tag, attrib=element.attrib, nsmap=element.nsmap)
    new_element.text = element.text
    new_element.tail = element.tail

    # Recursively copy child elements
    for child in element:
        new_child = manual_deepcopy(child)
        new_element.append(new_child)

    return new_element


def generate_nsid(nsid_list):
    new_nsid = "{:08X}".format(random.randint(0, 0xFFFFFFFF))
    while new_nsid in nsid_list:
        new_nsid = "{:08X}".format(random.randint(0, 0xFFFFFFFF))
    return new_nsid



def duplicate_numbering_format(doc):
    global root
    if root is None:
        # Access numbering part (numbering.xml)
        numbering_part = doc.part.numbering_part
        numbering_xml = etree.tostring(numbering_part._element, pretty_print=True).decode()
        # Parse numbering XML with lxml
        root = etree.fromstring(numbering_xml)

    # Find the 'w:numbering' element
    numbering = root.xpath('//w:numbering',
                               namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    highest_w_element = None
    if numbering is not None:
        w_elements = numbering[0].xpath('./w:num', namespaces={
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if w_elements:
            highest_w_element = max(int(e.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')) for e in w_elements)
        new_numId = highest_w_element + 1
        abstract_elements = numbering[0].xpath('./w:abstractNum', namespaces={
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if abstract_elements:
            highest_abstract_element = max(int(e.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')) for e in abstract_elements)
            nsid_list = [e.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') for e in numbering[0].findall('.//w:abstractNum/w:nsid', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})]

        if highest_abstract_element is not None:
            #print(f"Maximum abstract is {highest_abstract_element}")
            abstract_copy_elements = numbering[0].xpath(f'.//w:abstractNum[@w:abstractNumId="{highest_abstract_element}"]', namespaces={
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            new_abstract_element = manual_deepcopy(abstract_copy_elements[0])
            new_abstract_num = highest_abstract_element + 1
            new_abstract_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId', str(new_abstract_num))
            abstract_nsid = new_abstract_element.find(
                './{http://schemas.openxmlformats.org/wordprocessingml/2006/main}nsid')
            abstract_nsid.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(generate_nsid(nsid_list)))
            parent = abstract_copy_elements[0].getparent()
            pos_to_place = parent.index(abstract_copy_elements[0])
            numbering[0].insert(pos_to_place + 1, new_abstract_element)
        if highest_w_element is not None:
            maximum_numID = highest_w_element
            #print(f"Maximum W xml is {maximum_numID}")
            #print(f"Duplicating xml numbering for {new_numId}") #always use 1 as this is the right one
            w_copy_elements = numbering[0].xpath(f'.//w:num[@w:numId="1"]', namespaces={
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            new_numId = maximum_numID + 1
            if w_copy_elements:
                copy_w = w_copy_elements[0]
                # If the target 'w:abstractNum' exists, find the 'w:lvl' elements to duplicate
                if copy_w is not None:
                    # Create a new 'w:lvl' as a clone of the 'w:lvl' with '@w:ilvl=1'
                    # new_element = etree.Element(abstract_num.tag, nsmap=abstract_num.nsmap)
                    new_element = manual_deepcopy(copy_w)
                    # Update the 'numId' value
                    new_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId', str(new_numId))

                    # this sets a pseudo-random durable ID which is unique for the list but this does not seem to be necessary anyway
                    # new_element.set('{http://schemas.microsoft.com/office/word/2016/wordml/cid}durableId',
                    #                 str(generate_pseudo_random_durableId(new_numId)))
                    abstract_num_id = new_element.find('./{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                    abstract_num_id.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(new_abstract_num))
                    # # Append the new 'w:lvl' to 'w:numbering'
                    numbering[0].append(new_element)

    return new_numId


# Function to add a paragraph with numbering
def add_numbered_paragraph(doc,text, new_list=False, red=False):
    global new_number
    #new_list = False
    p = doc.add_paragraph(style='List Paragraph letter')
    if new_list:
        #print(f"Starting new list {new_number}")
        new_number = duplicate_numbering_format(doc)
        #print(f"Starting at {new_number}")
        restart_numbering(p, 0, new_number)
    else:
        #print(f"setting to {new_number}")
        restart_numbering(p, 0, new_number)
    run = p.add_run(text)
    if red:
        run.font.color.rgb = RGBColor(255, 0, 0)


def update_docx_numbering(doc):
    global root
    global new_number
    if root is not None:
        # new_xml = etree.tostring(root, pretty_print=True)
        numbering_part = doc.part.numbering_part
        # numbering_part._element.clear()
        # numbering_part._element.append(etree.fromstring(new_xml))
        numbering_part._element = root
        #reset these for the next document
        root = None
    new_number = 0


# Function to restart numbering
def restart_numbering(paragraph, ilvl_val=0, numId_val=3):
    ilvl_val = str(ilvl_val)
    numId_val = str(numId_val)
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    ilvl = numPr.get_or_add_ilvl()
    ilvl.set(qn('w:val'), ilvl_val)  # Set ilvl value
    numId = numPr.get_or_add_numId()
    numId.set(qn('w:val'), numId_val)  # Set numId value
    #this is not actually needed and has no effect
    # restart = OxmlElement('w:isRestart')
    # restart.set(qn('w:val'), '1')
    # numPr.insert(1, restart)  # Insert


def find_image_by_xid(xid, image_dir):
    for root, dirs, files in os.walk(image_dir):
        for file in files:
            if xid in file:
                return os.path.join(root, file)
    return None


def set_table_borders(table):
    table.style = 'Table Grid'


def get_text_width(document):
    """
    Returns the text width in mm.
    """
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 36000


def add_html_to_word(doc, html, colour, image_dir):
    soup = BeautifulSoup(html, 'html.parser')
    process_elements(doc, soup, colour, image_dir, doc.add_paragraph())


def process_elements(doc, elements, colour, image_dir, paragraph):
    for element in elements:
        if isinstance(element, NavigableString):
            handle_paragraph(doc, element, colour, paragraph)
            #print("found nav string")
        elif element.name == 'p':
            paragraph = doc.add_paragraph()
            process_elements(doc, element.contents, colour, image_dir, paragraph)
        elif element.name == 'table':
            handle_table(doc, element, colour)
            paragraph = doc.add_paragraph()
        elif element.name == 'img':
            handle_image(doc, element, image_dir)
        elif element.name == 'div':
            # Recursively process elements within the div
            paragraph = process_elements(doc, element.contents, colour, image_dir,
                                         paragraph)
        elif element.name == 'span' or element.name == 'i':
            paragraph = process_elements(doc, element.contents, colour, image_dir,
                                         paragraph)
        elif element.name == 'strong' or element.name == 'b' or element.name == 'em':
            text_content = element.get_text()
            if text_content.strip():  # Check if the text is not just whitespace
                run = paragraph.add_run(text_content)
                run.bold = True
                if colour:
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)
        elif element.name.startswith('h') and element.name[1:].isdigit():
            level = int(element.name[1:]) - 1
            text = element.get_text()
            doc.add_heading(text, level=level)
            paragraph = doc.add_paragraph()
        elif element.name == 'sup':
            text_content = element.get_text()
            if text_content.strip():  # Check if the text is not just whitespace
                run = paragraph.add_run(text_content)
                run.vertAlign = 'superscript'
                if colour:
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)
        elif element.name == 'sub':
            text_content = element.get_text()
            if text_content.strip():  # Check if the text is not just whitespace
                run = paragraph.add_run(text_content)
                run.vertAlign = 'subscript'
                if colour:
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)
        elif element.name == 'a':
            run = paragraph.add_run(element.text)
            if colour:
                font = run.font
                font.color.rgb = RGBColor(255, 0, 0)
        elif element.name == 'br':
            if not element.contents and not element.text:
                run= paragraph.add_run()
                run.add_break(WD_BREAK.LINE)
            else:
                paragraph = doc.add_paragraph()
                process_elements(doc, element.contents, colour, image_dir, paragraph)
        elif element.name == 'hr':
            paragraph.add_run("----------------------")
        else:
            print(f"Unhandled html element name: {element.name} {str(element)}")
            text_content = element.get_text()
            if text_content.strip():  # Check if the text is not just whitespace
                run = paragraph.add_run(text_content)
                if colour:
                    font = run.font
                    font.color.rgb = RGBColor(255, 0, 0)
    return paragraph


def handle_paragraph(doc, element, colour, paragraph):
    text_content = element.get_text()
    if text_content.strip():  # Check if the text is not just whitespace
        run = paragraph.add_run(text_content)
        if colour:
            font = run.font
            font.color.rgb = RGBColor(255, 0, 0)


# def handle_table(doc, element, colour):
#     table_data = []
#     for row in element.find_all('tr'):
#         row_data = []
#         for cell in row.find_all(['td', 'th']):
#             row_data.append(cell.text)
#         table_data.append(row_data)
#
#     # Check if table_data is empty
#     if not table_data:
#         return
#
#     # Find the maximum number of columns
#     col_count = max(len(row) for row in table_data)
#     row_count = len(table_data)
#
#     # Create the table
#     table = doc.add_table(rows=row_count, cols=col_count)
#     set_table_borders(table)
#
#     # Populate the table
#     for i, row_data in enumerate(table_data):
#         for j, cell_data in enumerate(row_data):
#             cell = table.cell(i, j)
#             cell.text = cell_data
#             if colour:
#                 for paragraph in cell.paragraphs:
#                     for run in paragraph.runs:
#                         run.font.color.rgb = RGBColor(255, 0, 0)


def handle_table(doc, element, colour):
    table_data = []
    for row in element.find_all('tr'):
        row_data = []
        for cell in row.find_all(['td', 'th']):
            text = cell.text
            colspan = int(cell.get('colspan', 1))
            row_data.append((text, colspan))
        table_data.append(row_data)

    if not table_data:
        return

    col_count = max(sum(cell[1] for cell in row) for row in table_data)
    row_count = len(table_data)

    table = doc.add_table(rows=row_count, cols=col_count)
    set_table_borders(table)

    for i, row_data in enumerate(table_data):
        j = 0
        for cell_text, cell_colspan in row_data:
            # Check if j exceeds the maximum number of columns
            if j >= col_count:
                print(f"Warning: Skipping cell at row {i}, column {j} due to incorrect colspan in table.")
                break

            try:
                cell = table.cell(i, j)
            except IndexError:
                print(f"Error: Could not access cell at row {i}, column {j} incorrect or unsupported table format.")
                break

            cell.text = cell_text
            if colour:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)

            j += cell_colspan


def handle_image(doc, element, image_dir):
    img_xid = re.search(r'/([^/]+)$', element.get('src',''))
    if img_xid:
        img_path = find_image_by_xid(img_xid.group(1), image_dir)
        if img_path:
            doc.add_picture(img_path, width=Mm(get_text_width(doc)))