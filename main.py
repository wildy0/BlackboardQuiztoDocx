# Convert Blackboard LMS Quiz ZIP to word.  Select the zip file, output files will be created for quiz and banks in
# a subdirectory at the source file location.  Please rename the zip to a short file length as long filenames
# can cause issues
# Tested on Windows for a variety of quiz types including MCQ and essay answer quiz.
# Created by Dr Tim Wilding,  2023
# Copyright (c) 2023, Dr Tim Wilding
# All rights reserved.
#
# This source code is licensed under the BSD-style license found in the
# LICENSE file in the root directory of this source tree.
import platform
import tkinter
from tkinter import filedialog
import os
import zipfile
from xml.etree import ElementTree as ET
from html import unescape
import shutil
from bs4 import BeautifulSoup, MarkupResemblesLocatorWarning
from docx.shared import RGBColor
import re
#import pypandoc
from docx import Document
import warnings
from htmltodocx import find_image_by_xid, add_html_to_word, set_table_borders, add_numbered_paragraph, update_docx_numbering
warnings.filterwarnings("ignore", category=MarkupResemblesLocatorWarning)


# Function to unescape HTML content and convert to plain text
def unescape_html(html_content):
    return unescape(html_content)


# Function to convert HTML content to plain text
def html_to_plain_text(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup.get_text()

# Function to find image file by its xid markup code in the image directory


def find_response_blocks(xml_content):
    response_blocks = []
    for flow in xml_content.findall('.//solutionmaterial'):
        content_elem = flow.find('.//mat_formattedtext')
        if content_elem is not None:
            response_blocks.append(content_elem.text)
    return response_blocks


def find_question_blocks(xml_content):
    response_blocks = []
    for flow in xml_content.findall('.//flow'):
        if 'class' in flow.attrib and flow.attrib['class'] == 'QUESTION_BLOCK':
            content_elem = flow.find('.//mat_formattedtext')
            if content_elem is not None:
                response_blocks.append(content_elem.text)
    return response_blocks


def find_correct_label(item):
    correct_label = []
    for cond in item.findall('.//respcondition'):
        if cond.attrib.get('title') == "correct":
            content_elem = cond.find('.//varequal')
            if content_elem is not None:
                correct_label.append(content_elem.text)
    return correct_label


def parse_question_banks(xml_content):
    banks = []
    root = ET.fromstring(xml_content)

    for item in root.findall(".//selection_ordering"):
        for itemz in item.findall(".//selection"):
            num = itemz.find('.//selection_number')
            source = itemz.find('.//sourcebank_ref')
            if num is not None and source is not None:
                banks.append([num.text, source.text])
    return banks


# Function to parse the XML content and extract questions
def parse_questions_advanced(xml_content):
    questions = []
    root = ET.fromstring(xml_content)

    for item in root.findall(".//item"):

        # flow class="RESPONSE_BLOCK" and flow class="QUESTION_BLOCK"

        question = {}

        question_block = find_question_blocks(item)
        answer_block = find_response_blocks(item)
        correct_label = find_correct_label(item)

        question['text'] = ""
        question['answer'] = ""

        for text in question_block:
            question['text'] = question['text'] + text


        first_block = True
        for text in answer_block:
            if text is not None:
                if first_block:
                    question['answer'] += text
                    # first block no need to check if contains a paragraph, just need that to separate subsequent ones
                else:
                    soup = BeautifulSoup(text, 'html.parser')
                    if len(soup.find_all('p')) == 0: #check if there is a <p>, if not add one to separate comments
                        question['answer'] += "<p>" + text + "</p>"
                    else:
                        question['answer'] += text
            first_block = False

        choices = []
        labels = []
        for response_label in item.findall(".//response_label"):
            label = response_label.attrib.get('ident', None)
            choice_elem = response_label.find(".//mat_formattedtext")
            if choice_elem is not None:
                choices.append(unescape_html(choice_elem.text))
            if label is not None:
                labels.append(label)

        question['choices'] = choices
        question['labels'] = labels
        question['correct'] = correct_label

        questions.append(question)

    return questions


def append_docx(doc,append_doc_path):
    append_doc = Document(append_doc_path)
    for element in append_doc.element.body:
        doc.element.body.append(element)


def add_html_to_word_pandoc(html, colour, image_dir):
    if html != "":
        #first find all img tags and locate the image source
        soup = BeautifulSoup(html, 'html.parser')
        for img_tag in soup.find_all('img'):
            img_src = img_tag['src']
            img_xid = re.search(r'/([^/]+)$', img_src)
            if img_xid:
                img_path = find_image_by_xid(img_xid.group(1), image_dir)
                if img_path:
                    new_src = find_image_by_xid(img_xid.group(1), image_dir)
                    img_tag['src'] = new_src
        #update the html with the new html containing the new correct image source locations
        if colour:
            html = "<div style=\"color: red;\">" + str(soup) + "</div>"
        else:
            html = str(soup)
    return html


# this is for a pandoc convert which does not support text colours easily,
# needs work manual conversion to docx works ok for now
def convert_quiz_to_word_advanced_pandoc(dat_content, word_file_path, image_dir=None):
    #convert_quiz_to_word_advanced(dat_content,word_file_path,image_dir)
    #doc = Document()
    html = """
<!DOCTYPE html>
<html>
<head>
  <style>
    ol.alphabetical {
      list-style-type: none;
      counter-reset: list-counter;
    }
    ol.alphabetical li {
      counter-increment: list-counter;
    }
    ol.alphabetical li:before {
      content: counter(list-counter, lower-alpha) ". ";
    }
  </style>
</head>
<body> 
    """
    questions = parse_questions_advanced(dat_content)

    if not bool(questions):
        html += 'Test file for question banks'
        banks = parse_question_banks(dat_content)
        html += '<table>'
        for b in banks:
            number, bank = b
            html += (f"<tr><td>{number}</td><td>{bank}</td></tr>")
        html += '</table>'
    else:
        for i, question in enumerate(questions):
            html += f"<p>Question: {i+1}</ br>"
            html += add_html_to_word_pandoc(question.get('text', 'N/A'), False, image_dir)
            correct_label = question.get('correct')
            labels = question.get('labels')
            answer_text = ""
            html += "<ol class=\"alphabetical\">"
            for j, (label, choice) in enumerate(zip(labels, question.get('choices', []))):
                #para = doc.add_paragraph()#style="List Bullet")
                #para.style.number_format = 'a'
                choice_text = html_to_plain_text(choice)
                #run = para.add_run(f"{chr(65+j)}) {choice_text}")
                #html += "<ol>"
                if label in correct_label:
                    #html += f"<li><span style=\"color:red;\">{chr(65+j)}{choice_text}</span></li>"
                    html += f"<li>{choice_text}</li>"
                    #answer_text = f"Answer: {chr(65+j)}"
                    answer_text = f"Answer:{chr(65+j)}"
                else:
                    #html += f"<li>{chr(65+j)}{choice_text}</li>"
                    html += f"<li>{choice_text}</li>"
            html += "</ol></p>"
            if answer_text != "":
                html += answer_text
    html += add_html_to_word_pandoc(question.get('answer', 'N/A'), True, image_dir)
    html += "</body></html>"
    pypandoc.convert_text(html, 'docx', format='html', outputfile=word_file_path)


# Function to convert quiz data to Word document
def convert_quiz_to_word_advanced(dat_content, word_file_path, bank_names, image_dir=None, lams=False):
    doc = Document('template.docx')
    questions = parse_questions_advanced(dat_content)
    is_mcq = False
    if not bool(questions):
        doc.add_paragraph('Test file for question banks')
        banks = parse_question_banks(dat_content)

        table = doc.add_table(rows=len(banks)+2, cols=2)
        set_table_borders(table)
        cell = table.cell(0, 0)
        cell.text = 'Bank Name'
        cell = table.cell(0, 1)
        cell.text = "Number of questions taken from bank"
        total_number = 0
        for i, b in enumerate(banks):
            number, bank = b
            #doc.add_paragraph(f"{number} {bank_names.get(bank,'bank name unknown')}")
            cell = table.cell(i+1, 0)
            cell.text = bank_names.get(bank,'bank name unknown')
            cell = table.cell(i+1, 1)
            cell.text = number
            total_number += int(number)

        cell = table.cell(len(banks)+1, 0)
        cell.text = 'Total'
        cell = table.cell(len(banks)+1, 1)
        cell.text = str(total_number)

    else:
        for i, question in enumerate(questions):
            if lams:
                doc.add_paragraph(f"Question: {i+1}")
            else:
                doc.add_heading(f"Question: {i+1}", level=2)
            add_html_to_word(doc, question.get('text', 'N/A'), False, image_dir)
            correct_label = question.get('correct')
            labels = question.get('labels')
            answer_text = ""
            new = True
            red = False
            for j, (label, choice) in enumerate(zip(labels, question.get('choices', []))):
                #para = doc.add_paragraph()#style="List Bullet")
                #para.style.number_format = 'a'
                is_mcq = True
                choice_text = html_to_plain_text(choice)
                if j == 26:
                    #here we have a problem because if j=26 then we have more than a-z and the code will create options
                    #with characters from { to wierd control codes
                    answer_text = "Answer: error" #set the answer to error to show it went wrong, but this is
                    # not likely to ever happen
                    break #break the loop so we don't continue past z
                option_letter = chr(97 + j)
                if label in correct_label:
                    if lams:
                        if answer_text == "":
                            answer_text = f"Answer: {option_letter}"
                        else:
                            answer_text += f",{option_letter}"
                        red = False
                    else:
                        red = True
                else:
                    red = False

                if not lams:
                    add_numbered_paragraph(doc, choice_text, new_list=new, red=red)
                else:
                    #lams is fussy with lists and so we need it like this
                    doc.add_paragraph(f"{option_letter}) {choice_text}")
                if new:
                    new = False
                #run = para.add_run(f"{chr(65+j)}) {choice_text}")
                # run = para.add_run(f"{chr(97+j)}) {choice_text}")
                # if label == correct_label:
                #     run.font.color.rgb = RGBColor(255, 0, 0)
                    #answer_text = f"Answer: {chr(65+j)}"
            if answer_text != "":
                doc.add_paragraph(answer_text)

            if not lams:
                add_html_to_word(doc, question.get('answer', 'N/A'), True, image_dir)

    update_docx_numbering(doc)
    doc.save(word_file_path)
    return is_mcq


def find_assessment_filenames(xml_content):
    root = ET.fromstring(xml_content)
    filenames = []
    names = []
    for resource in root.findall('.//resource'):
        resource_type = resource.attrib.get('type', '')
        if resource_type.startswith('assessment/'):
            bb_file = resource.attrib.get('{http://www.blackboard.com/content-packaging/}file', None)
            if bb_file is not None:
                filenames.append(bb_file)
            bb_title = resource.attrib.get('{http://www.blackboard.com/content-packaging/}title', None)
            if bb_title is not None:
                if bb_title.startswith("single.qti.export"):
                    final_title = bb_title[len("single.qti.export.referenced.canvas.name.prefix "):]
                else:
                    final_title = bb_title
                final_title = re.sub(r'[^\w\-]', '_', final_title)
                names.append(final_title)
    return filenames, names


def sanitize_filename(filename):
    """
    Remove invalid characters from filename string.
    """
    sanitized = re.sub(r'[\\/:*?"<>|]', '', filename)
    sanitized = sanitized.replace(' ', '_')
    sanitized = sanitized.replace('-', '_')
    return sanitized


# Main function to process the Blackboard zip export
def process_blackboard_zip(zip_file_path):

    base_path = os.path.dirname(zip_file_path)
    #os.chdir(base_path)

    # Create directories for extraction and output
    base_name = os.path.basename(zip_file_path)
    extracted_folder_path = os.path.join(base_path,sanitize_filename(os.path.splitext(base_name)[0] + '_zip'))
    output_word_file_path = os.path.join(base_path,sanitize_filename(os.path.splitext(base_name)[0] + '_docx'))

    if not os.path.exists(extracted_folder_path):
        os.makedirs(extracted_folder_path)

    if not os.path.exists(output_word_file_path):
        os.makedirs(output_word_file_path)

    # Unzip the Blackboard export
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(extracted_folder_path)

    manifest_file_path = os.path.join(extracted_folder_path, 'imsmanifest.xml')
    with open(manifest_file_path, 'r') as f:
        xml_content = f.read()

    dat_files, dat_names = find_assessment_filenames(xml_content)
    image_dir = os.path.join(extracted_folder_path, 'csfiles/home_dir')

    # parse to get bank_names
    bank_names = {}
    for dat_file, dat_name in zip(dat_files, dat_names):
        bank_names[dat_file.split(".")[0]] = dat_name

    # Iterate through each .dat file to convert its content to the Word document
    for dat_file, dat_name in zip(dat_files, dat_names):
        filename = os.path.join(extracted_folder_path, dat_file)
        with open(filename, 'r', encoding='utf-8', errors='ignore') as f:
            dat_content = f.read()
        # Convert the quiz to a Word document
        #     try:
        #         pandoc_version = pypandoc.get_pandoc_version()
        #         print(f"Pandoc version {pandoc_version} is available.")
        #         convert_quiz_to_word_advanced_pandoc(dat_content, os.path.join(output_word_file_path, f"{os.path.splitext(os.path.basename(dat_file))[0]}_{dat_name}.docx"), image_dir=image_dir)
        #     except OSError:
        #         print("Pandoc is not available. Using fallback.")
        #         convert_quiz_to_word_advanced(dat_content, os.path.join(output_word_file_path, f"{os.path.splitext(os.path.basename(dat_file))[0]}_{dat_name}.docx"), image_dir=image_dir)
            save_name = f"{os.path.splitext(os.path.basename(dat_file))[0]}" if dat_name == '' else dat_name
            print(f"Parsing file {dat_file} to word output {save_name}.docx")
            if convert_quiz_to_word_advanced(dat_content, os.path.join(output_word_file_path, f"{save_name}.docx"),
                                          bank_names, image_dir=image_dir,lams=False):
                convert_quiz_to_word_advanced(dat_content, os.path.join(output_word_file_path, f"{save_name}_lams.docx"),
                                              bank_names, image_dir=image_dir, lams=True)

    shutil.rmtree(extracted_folder_path)
    print(f"Output files are located at {output_word_file_path}")
    input("Press/Enter any key to continue...")



if __name__ == '__main__':
    root = tkinter.Tk()
    root.withdraw()
    # root.update()
    filename = filedialog.askopenfilename(title="Blackboard ZIP",
                                              filetypes=[("Blackboard Quiz  (zip)", ".zip")]
                                              )
    if platform.system() == "Windows":
        filename = filename.replace("/", "\\")
    print("Reading file %s" % filename)
    process_blackboard_zip(filename)

